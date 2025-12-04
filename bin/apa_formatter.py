#!/usr/bin/env python3
"""
APA Formatter - Trinity's Academic Document Formatter
A simple desktop tool to convert markdown/text to APA-formatted Word documents

Usage:
    python apa_formatter.py                    # Launch GUI
    python apa_formatter.py input.md           # Convert file directly
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import sys
import os


class APAFormatter:
    """Core formatting logic for APA documents"""
    
    @staticmethod
    def create_apa_document(title, author, institution, content, output_path, 
                            include_abstract=False, abstract_text=''):
        """
        Create an APA 7th edition formatted Word document
        
        Args:
            title: Document title
            author: Author name
            institution: Affiliated institution
            content: Markdown or plain text content
            output_path: Path to save .docx file
            include_abstract: Whether to include an abstract page
            abstract_text: Abstract content if included
        """
        doc = Document()
        
        # ==== APA FORMATTING STANDARDS ====
        # Font: Times New Roman 12pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Margins: 1 inch all sides
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Line spacing: Double-spaced
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 2.0
        
        # ==== TITLE PAGE (APA 7th Edition) ====
        # Running head (shortened title, max 50 chars)
        running_head = doc.add_paragraph()
        running_head.text = title.upper()[:50]
        running_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = running_head.runs[0]
        run.font.size = Pt(12)
        
        # Vertical centering (approx)
        for _ in range(8):
            doc.add_paragraph()
        
        # Title (bold, centered)
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Author
        author_para = doc.add_paragraph(author)
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Institution
        inst_para = doc.add_paragraph(institution)
        inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Course/Assignment info (optional)
        doc.add_paragraph()
        
        # Due date
        date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # ==== ABSTRACT (Optional) ====
        if include_abstract:
            abstract_heading = doc.add_paragraph('Abstract')
            abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            abstract_run = abstract_heading.runs[0]
            abstract_run.bold = True
            
            doc.add_paragraph(abstract_text)
            doc.add_paragraph()
            
            # Keywords
            keywords_para = doc.add_paragraph()
            keywords_run = keywords_para.add_run('Keywords: ')
            keywords_run.italic = True
            keywords_para.add_run('[Add keywords here]')
            
            doc.add_page_break()
        
        # ==== CONTENT ====
        # Parse markdown content
        APAFormatter._parse_and_add_content(doc, content)
        
        # Save
        doc.save(output_path)
        return True
    
    @staticmethod
    def _parse_and_add_content(doc, content):
        """Parse markdown and add formatted content to document"""
        lines = content.split('\n')
        in_code_block = False
        in_list = False
        
        for line in lines:
            # Skip main title (already on title page)
            if line.startswith('# '):
                continue
            
            # Code blocks
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
            
            if in_code_block:
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.0  # Single-spaced code
                if p.runs:
                    p.runs[0].font.name = 'Courier New'
                    p.runs[0].font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.5)
                continue
            
            # Headings
            if line.startswith('## '):
                in_list = False
                p = doc.add_paragraph()
                run = p.add_run(line[3:].strip())
                run.bold = True
                run.font.size = Pt(14)
            elif line.startswith('### '):
                in_list = False
                p = doc.add_paragraph()
                run = p.add_run(line[4:].strip())
                run.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.left_indent = Inches(0.25)
            elif line.startswith('#### '):
                in_list = False
                p = doc.add_paragraph()
                run = p.add_run(line[5:].strip())
                run.italic = True
                p.paragraph_format.left_indent = Inches(0.5)
            
            # Lists
            elif line.strip().startswith(('- ', '* ')):
                in_list = True
                text = line.strip()[2:]
                doc.add_paragraph(text, style='List Bullet')
            elif re.match(r'^\d+\. ', line.strip()):
                in_list = True
                text = re.sub(r'^\d+\. ', '', line.strip())
                doc.add_paragraph(text, style='List Number')
            
            # Horizontal rules
            elif line.strip() == '---':
                in_list = False
                doc.add_paragraph('_' * 60)
            
            # Empty lines
            elif not line.strip():
                if in_list:
                    in_list = False
                else:
                    doc.add_paragraph()
            
            # Regular paragraphs
            else:
                in_list = False
                # Handle inline formatting (**bold**, *italic*, `code`)
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    else:
                        p.add_run(part)


class APAFormatterGUI:
    """GUI for APA Formatter"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("APA Formatter - Trinity's Academic Document Tool")
        self.root.geometry("800x700")
        
        # Configure style
        style = ttk.Style()
        style.theme_use('clam')
        
        self.create_widgets()
    
    def create_widgets(self):
        # Main container
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Title
        title_label = ttk.Label(main_frame, text="APA 7th Edition Document Formatter", 
                                font=('Arial', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=2, pady=10)
        
        # === METADATA SECTION ===
        metadata_frame = ttk.LabelFrame(main_frame, text="Document Metadata", padding="10")
        metadata_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # Title
        ttk.Label(metadata_frame, text="Title:").grid(row=0, column=0, sticky=tk.W, pady=2)
        self.title_entry = ttk.Entry(metadata_frame, width=60)
        self.title_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=2)
        
        # Author
        ttk.Label(metadata_frame, text="Author:").grid(row=1, column=0, sticky=tk.W, pady=2)
        self.author_entry = ttk.Entry(metadata_frame, width=60)
        self.author_entry.insert(0, "Trinity")
        self.author_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=2)
        
        # Institution
        ttk.Label(metadata_frame, text="Institution:").grid(row=2, column=0, sticky=tk.W, pady=2)
        self.institution_entry = ttk.Entry(metadata_frame, width=60)
        self.institution_entry.insert(0, "Purdue University")
        self.institution_entry.grid(row=2, column=1, sticky=(tk.W, tk.E), pady=2)
        
        # Abstract checkbox
        self.include_abstract = tk.BooleanVar()
        abstract_check = ttk.Checkbutton(metadata_frame, text="Include Abstract Page", 
                                         variable=self.include_abstract,
                                         command=self.toggle_abstract)
        abstract_check.grid(row=3, column=0, columnspan=2, sticky=tk.W, pady=5)
        
        # === CONTENT SECTION ===
        content_frame = ttk.LabelFrame(main_frame, text="Content (Markdown or Plain Text)", padding="10")
        content_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Content text area
        self.content_text = scrolledtext.ScrolledText(content_frame, width=70, height=20, wrap=tk.WORD)
        self.content_text.grid(row=0, column=0, pady=5)
        
        # Placeholder text
        placeholder = """# Your Document Title

## Introduction
Start writing your content here. You can use markdown formatting:

- **Bold text** with double asterisks
- *Italic text* with single asterisks
- `Code` with backticks

## Headings
Use ## for H2, ### for H3, etc.

## Lists
1. Numbered lists
2. Like this

- Or bullet points
- Like this
"""
        self.content_text.insert('1.0', placeholder)
        
        # === ACTION BUTTONS ===
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=10)
        
        # Load file button
        load_btn = ttk.Button(button_frame, text="Load Markdown File", command=self.load_file)
        load_btn.grid(row=0, column=0, padx=5)
        
        # Convert button
        convert_btn = ttk.Button(button_frame, text="Convert to APA .docx", 
                                 command=self.convert_to_docx, style='Accent.TButton')
        convert_btn.grid(row=0, column=1, padx=5)
        
        # Clear button
        clear_btn = ttk.Button(button_frame, text="Clear All", command=self.clear_all)
        clear_btn.grid(row=0, column=2, padx=5)
        
        # === STATUS ===
        self.status_label = ttk.Label(main_frame, text="Ready", foreground="green")
        self.status_label.grid(row=4, column=0, columnspan=2, pady=5)
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(2, weight=1)
        content_frame.columnconfigure(0, weight=1)
        content_frame.rowconfigure(0, weight=1)
    
    def toggle_abstract(self):
        """Show/hide abstract input when checkbox toggled"""
        # In future version, add abstract text input
        pass
    
    def load_file(self):
        """Load markdown file into content area"""
        filename = filedialog.askopenfilename(
            title="Select Markdown File",
            filetypes=[("Markdown files", "*.md"), ("Text files", "*.txt"), ("All files", "*.*")]
        )
        if filename:
            try:
                with open(filename, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.content_text.delete('1.0', tk.END)
                self.content_text.insert('1.0', content)
                self.status_label.config(text=f"Loaded: {os.path.basename(filename)}", foreground="blue")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load file:\n{str(e)}")
    
    def convert_to_docx(self):
        """Convert content to APA-formatted .docx"""
        # Get metadata
        title = self.title_entry.get().strip()
        author = self.author_entry.get().strip()
        institution = self.institution_entry.get().strip()
        content = self.content_text.get('1.0', tk.END)
        
        # Validation
        if not title:
            messagebox.showwarning("Missing Information", "Please enter a document title.")
            return
        if not author:
            messagebox.showwarning("Missing Information", "Please enter an author name.")
            return
        
        # Ask where to save
        output_path = filedialog.asksaveasfilename(
            title="Save APA Document",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")],
            initialfile=f"{title.replace(' ', '_')}.docx"
        )
        
        if not output_path:
            return
        
        try:
            # Convert
            self.status_label.config(text="Converting...", foreground="orange")
            self.root.update()
            
            APAFormatter.create_apa_document(
                title=title,
                author=author,
                institution=institution,
                content=content,
                output_path=output_path,
                include_abstract=self.include_abstract.get()
            )
            
            self.status_label.config(text=f"✓ Created: {os.path.basename(output_path)}", 
                                     foreground="green")
            
            # Ask to open
            if messagebox.askyesno("Success", "Document created! Open it now?"):
                os.startfile(output_path)
        
        except Exception as e:
            messagebox.showerror("Conversion Error", f"Failed to create document:\n{str(e)}")
            self.status_label.config(text="Error - see message", foreground="red")
    
    def clear_all(self):
        """Clear all fields"""
        if messagebox.askyesno("Clear All", "Clear all fields?"):
            self.title_entry.delete(0, tk.END)
            self.content_text.delete('1.0', tk.END)
            self.include_abstract.set(False)
            self.status_label.config(text="Cleared", foreground="gray")


def main():
    """Main entry point"""
    # Command-line mode
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        if not os.path.exists(input_file):
            print(f"Error: File not found: {input_file}")
            sys.exit(1)
        
        # Read file
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract title from first heading
        title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
        title = title_match.group(1) if title_match else "Document"
        
        # Output path
        output_path = input_file.rsplit('.', 1)[0] + '.docx'
        
        # Convert
        print(f"Converting {input_file} to APA format...")
        APAFormatter.create_apa_document(
            title=title,
            author="Trinity",
            institution="Purdue University",
            content=content,
            output_path=output_path
        )
        print(f"✓ Created: {output_path}")
        sys.exit(0)
    
    # GUI mode
    root = tk.Tk()
    app = APAFormatterGUI(root)
    root.mainloop()


if __name__ == '__main__':
    main()
